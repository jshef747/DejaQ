"""Fingerprint for uploaded FILES (PDF, DOCX, Markdown, text/code) — the file gate.

Read this before reaching for anything in image_text.py or image_fingerprint.py:
the file gate is deliberately NOT built like the image gate, and the difference is
not a shortcut.

The image gate needs CLIP, dHash, tile variety, OCR confidence floors and a swept
token-overlap threshold because OCR output is *noisy*: two reads of the same page
disagree on characters, so identity can only be approximated, and every constant
had to be measured over ~1.85M labelled pairs to find where approximation stops
being safe.

A PDF, DOCX, or text file hands us the text directly and deterministically. The
same bytes always extract the same characters. So identity here is EXACT — a
sha256 over the whitespace-normalised text — and the gate is string equality.
Two different documents cannot collide, so there is no false-merge rate, no
recall-vs-merge curve, nothing to sweep, and no eval harness. That is the entire
design, and it is why this module is short.

Whitespace is normalised (not stripped of content, not case-folded) so that the
same document re-saved by a different producer — which commonly reflows line
breaks and spacing while preserving the words — still hashes the same. Case is
left alone: the same file always extracts the same case, so folding it would buy
nothing and could only merge two documents that genuinely differ.

A file we cannot read — a scanned PDF with no text layer, an encrypted PDF, a
corrupt upload — returns ok=False and is then neither served nor stored, mirroring
the image gate's `ambiguous` class. Answering it still works; it just never
becomes a cache entry. Nothing here raises into a request.
"""
from __future__ import annotations

import hashlib
import io
import logging
import re
from dataclasses import dataclass

from app.config import CACHE_FILE_MIN_CHARS

logger = logging.getLogger("dejaq.services.file_text")

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_DOCX_SUFFIXES = (".docx",)

_WHITESPACE_RE = re.compile(r"\s+")

_pypdf_missing_warned = False
_docx_missing_warned = False


@dataclass(frozen=True)
class FileText:
    """The result of reading an uploaded file.

    `text` is the extracted content as-is (what gets sent to the model for
    Markdown); `sha` is the identity used by the gate. They are deliberately
    different: the model wants readable text, the gate wants a stable key.
    """

    kind: str          # "pdf" | "docx" | "markdown" | "" when the type is not supported
    text: str          # extracted content, un-normalised
    sha: str           # sha256 of the normalised text; "" when not identifiable
    char_count: int    # length of the normalised text
    ok: bool           # False -> never served, never stored
    reason: str        # why not ok, for the logs; "" when ok

    @property
    def cacheable(self) -> bool:
        return self.ok and bool(self.sha)

    @property
    def readable(self) -> bool:
        """True when there is any real text to answer from.

        Deliberately independent of `ok`/`cacheable`: `ok` also requires
        clearing CACHE_FILE_MIN_CHARS, a floor that exists only to keep a
        near-empty extraction from being hashed into a cache identity (see
        `_finalize`). A short-but-genuine document (a 4-line memo, a one-line
        receipt) extracts real, complete text below that floor - refusing to
        answer from it would be gating answering on a caching threshold, which
        it must never be. `char_count` (normalised length) is 0 for every
        genuinely unreadable case - unsupported type, decode/parse failure,
        scanned PDF with no text layer - so those still correctly return False.
        """
        return self.char_count > 0


def kind_for(data: bytes, mime: str | None, filename: str | None) -> str:
    """Which extractor handles this upload, or "" if we do not support it.

    PDF and DOCX are recognised by MIME or extension only, and checked before
    anything else: a `.docx` is a ZIP under the hood, so it would otherwise fall
    through to the text sniff below and be rejected (ZIP bytes are not valid
    UTF-8) or, worse, be accidentally readable as garbage text. `application/zip`
    is deliberately not in the DOCX MIME check — accepting it would make every
    ZIP archive look like a Word document; the extension is the trustworthy
    signal here, same as it is for Markdown.

    Everything else - Markdown, plain text, and every kind of source or config
    file - is decided by content, not a maintained extension or MIME list: if
    the bytes are a strict UTF-8 decode, it is text and reuses the existing
    "markdown" kind, because it is inlined into the prompt exactly the same
    way regardless of what produced it. This is what makes a `.py` file with an
    empty or wrong MIME behave identically to one with the "right" MIME - see
    file_gate.md for the inconsistency this replaced. It also means a `.md` or
    `.txt` file that happens not to be valid UTF-8 is correctly refused instead
    of silently decoded with replacement characters.
    """
    mime = (mime or "").split(";", 1)[0].strip().lower()
    name = (filename or "").strip().lower()
    if mime == PDF_MIME or name.endswith(".pdf"):
        return "pdf"
    if mime == DOCX_MIME or name.endswith(_DOCX_SUFFIXES):
        return "docx"
    if not data:
        return ""
    try:
        data.decode("utf-8", errors="strict")
    except UnicodeDecodeError:
        return ""
    return "markdown"


def _normalize(text: str) -> str:
    """Collapse every whitespace run to one space and trim.

    This is the only transformation between the extracted text and the hash, so
    it defines what "the same document" means. Keep it boring: anything cleverer
    here (stripping punctuation, folding case, dropping headers) would start
    merging documents that are not the same.
    """
    return _WHITESPACE_RE.sub(" ", text).strip()


def _finalize(kind: str, text: str, *, apply_floor: bool) -> FileText:
    """Hash the normalised text, or explain why there is nothing to hash.

    `apply_floor` gates CACHE_FILE_MIN_CHARS. It exists at all because
    extraction from PDF/DOCX can fail *silently* - a scanned page or a corrupt
    document yields near-nothing, and hashing "almost nothing" would give every
    unreadable document of that kind the same identity, a real cross-document
    leak. Text has no extraction step to go wrong: decoding it is exact, so a
    50-character script genuinely is 50 characters and its hash is a true
    identity, not a guess at one. The floor would only refuse legitimate short
    files there, so text/markdown skip it and require just one character.
    """
    normalized = _normalize(text)
    floor = CACHE_FILE_MIN_CHARS if apply_floor else 1
    if len(normalized) < floor:
        reason = (
            f"{len(normalized)} chars extracted (need >= {floor})"
            if apply_floor
            else "no text after normalisation"
        )
        return FileText(
            kind=kind,
            text=text,
            sha="",
            char_count=len(normalized),
            ok=False,
            reason=reason,
        )
    return FileText(
        kind=kind,
        text=text,
        sha=hashlib.sha256(normalized.encode("utf-8")).hexdigest(),
        char_count=len(normalized),
        ok=True,
        reason="",
    )


def _extract_pdf(data: bytes) -> str:
    """Every page's text, concatenated. Raises only what extract() catches."""
    from pypdf import PdfReader  # imported lazily so the import cost is paid once, on use

    reader = PdfReader(io.BytesIO(data))
    if reader.is_encrypted:
        # An empty-password decrypt covers the common "protected but not really"
        # case; anything else falls through to a short read and is refused.
        try:
            reader.decrypt("")
        except Exception:
            return ""
    # ponytail: reads every page. Bounded in practice by MAX_ATTACHMENT_BYTES,
    # and it runs on a threadpool. If very large PDFs ever become common, cap the
    # page count and hash a page-count-tagged prefix — do NOT hash a bare prefix,
    # or two documents sharing an opening chapter would merge.
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(data: bytes) -> str:
    """Paragraph and table text only. Raises only what extract() catches.

    Deliberately NOT extracted: headers, footers, footnotes, comments, tracked
    changes, document properties. Every one of those can change without the
    document meaningfully changing - someone adds a review comment, Word bumps
    a "last modified" property, a footer page number ticks over - and none of
    that is a reason for the cache to treat it as a different file. Hashing
    them in would make the cache miss on effectively every re-save, which looks
    exactly like a broken feature with no error to explain why.
    """
    import docx  # imported lazily so the import cost is paid once, on use

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def extract(data: bytes, mime: str | None = None, filename: str | None = None) -> FileText:
    """Read an uploaded file. Never raises — a failure yields ok=False."""
    global _pypdf_missing_warned, _docx_missing_warned

    kind = kind_for(data, mime, filename)
    if not kind:
        return FileText("", "", "", 0, ok=False, reason=f"unsupported type {mime!r}")

    if kind == "markdown":
        return _finalize(kind, data.decode("utf-8", errors="strict"), apply_floor=False)

    if kind == "docx":
        try:
            return _finalize(kind, _extract_docx(data), apply_floor=True)
        except ModuleNotFoundError:
            if not _docx_missing_warned:
                logger.warning(
                    "python-docx is not installed — DOCX uploads cannot be cached. "
                    "Install it (uv sync) to enable DOCX caching."
                )
                _docx_missing_warned = True
            return FileText(kind, "", "", 0, ok=False, reason="python-docx not installed")
        except Exception as exc:
            # Corrupt, encrypted, or otherwise unreadable. A miss, never a raise.
            logger.info("DOCX could not be read (%s); it will not be cached", type(exc).__name__)
            return FileText(kind, "", "", 0, ok=False, reason=f"unreadable ({type(exc).__name__})")

    try:
        return _finalize(kind, _extract_pdf(data), apply_floor=True)
    except ModuleNotFoundError:
        if not _pypdf_missing_warned:
            logger.warning(
                "pypdf is not installed — PDF uploads cannot be cached. "
                "Install it (uv sync) to enable PDF caching."
            )
            _pypdf_missing_warned = True
        return FileText(kind, "", "", 0, ok=False, reason="pypdf not installed")
    except Exception as exc:
        # Corrupt, encrypted, or otherwise unreadable. A miss, never a raise.
        logger.info("PDF could not be read (%s); it will not be cached", type(exc).__name__)
        return FileText(kind, "", "", 0, ok=False, reason=f"unreadable ({type(exc).__name__})")


def matches(new: FileText, stored_sha: str | None, stored_kind: str | None) -> bool:
    """Same file, exactly. No threshold, and there is not supposed to be one."""
    return bool(
        new.cacheable
        and stored_sha
        and new.sha == stored_sha
        and new.kind == (stored_kind or "")
    )


def _self_test() -> None:
    long_text = " ".join(f"word{i}" for i in range(200))
    assert len(_normalize(long_text)) >= CACHE_FILE_MIN_CHARS
    long_bytes = long_text.encode("utf-8")

    # Type routing: MIME or extension for PDF/DOCX, content for everything else.
    assert kind_for(b"", "application/pdf", None) == "pdf"
    assert kind_for(b"", None, "contract.PDF") == "pdf"
    assert kind_for(b"PK\x03\x04fake docx bytes", None, "report.DOCX") == "docx"
    assert kind_for(long_bytes, "text/plain", "notes.md") == "markdown"
    assert kind_for(long_bytes, "", "notes.md") == "markdown", "browsers often send no MIME for .md"
    assert kind_for(long_bytes, "text/x-python", "script.py") == "markdown", "code is text, MIME agnostic"
    assert kind_for(long_bytes, "video/mp2t", "main.ts") == "markdown", "browsers misreport .ts as video"
    assert kind_for(long_bytes, "application/json", "data.json") == "markdown"
    assert kind_for(b"\x89PNG\x0d\x0a\x1a\x0a\x00\x01", "image/png", "photo.png") == "", "not valid UTF-8"
    assert kind_for(b"", "", None) == "", "no MIME, no name, no bytes is unknown"

    a = extract(long_bytes, "text/markdown", "a.md")
    assert a.ok and a.kind == "markdown"

    # The whole point: same content -> same key, whatever the producer did to the
    # whitespace. Different content -> different key.
    reflowed = extract(long_text.replace(" ", "\n  ").encode(), "text/markdown", "b.md")
    assert reflowed.sha == a.sha, "re-flowed whitespace must not change identity"
    other = extract((long_text + " extra").encode(), "text/markdown", "c.md")
    assert other.sha != a.sha, "different content must not collide"

    # Same code file, wildly different (or absent) MIME -> identical acceptance.
    code = "def handler():\n    return 42\n" * 20
    by_correct_mime = extract(code.encode(), "text/x-python", "script.py")
    by_wrong_mime = extract(code.encode(), "video/mp2t", "script.py")
    by_no_mime = extract(code.encode(), "", "script.py")
    assert by_correct_mime.ok and by_wrong_mime.ok and by_no_mime.ok
    assert by_correct_mime.sha == by_wrong_mime.sha == by_no_mime.sha

    # Text/markdown has no floor: a short file is still cached, non-empty only.
    short = extract(b"hi", "text/markdown", "t.md")
    assert short.ok and short.cacheable, "a short text file has an exact identity, no floor applies"
    empty = extract(b"   \n\t  ", "text/markdown", "blank.md")
    assert not empty.ok, "whitespace-only normalises to nothing"

    # Not valid UTF-8 -> unsupported, not silently decoded.
    not_utf8 = extract(b"\xff\xfe\x00binary garbage", "text/plain", "weird.py")
    assert not not_utf8.ok and not_utf8.kind == ""

    # A short-but-genuine PDF/DOCX: uncacheable (under the floor) but readable
    # (real, complete text) - answering must not be gated on the caching floor.
    short_pdf = FileText("pdf", "Invoice total: 42 dollars.", "", 26, ok=False, reason="26 chars extracted (need >= 200)")
    assert not short_pdf.ok and short_pdf.readable, "short genuine extraction must still be answerable"

    # Unreadable PDF bytes are a miss, not an exception.
    broken = extract(b"%PDF-1.4 not really a pdf", "application/pdf", "x.pdf")
    assert not broken.ok and broken.kind == "pdf"
    assert not broken.readable, "a genuinely unreadable file must not be answerable either"

    # Unreadable DOCX bytes are a miss, not an exception, and never sniffed as text.
    broken_docx = extract(b"PK\x03\x04 not really a docx", DOCX_MIME, "x.docx")
    assert not broken_docx.ok and broken_docx.kind == "docx"

    # Unsupported types never reach a gate.
    assert not extract(b"\x89PNG", "image/png", "p.png").ok

    # The gate itself.
    assert matches(a, a.sha, "markdown")
    assert not matches(a, other.sha, "markdown"), "different files must not match"
    assert not matches(a, a.sha, "pdf"), "kinds must not mix"
    assert not matches(a, None, None), "an entry with no file is not a match"
    assert not matches(broken, broken.sha, "pdf"), "an unidentifiable file never matches"
    print("self-test passed")


if __name__ == "__main__":
    _self_test()
